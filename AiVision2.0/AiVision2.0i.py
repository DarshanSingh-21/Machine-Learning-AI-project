import os
import io
import sys
import time
import re
from pathlib import Path
from dotenv import load_dotenv

# SDK Engine Imports
from google import genai
from google.genai import types
import google.genai.errors
from groq import Groq

# Local PDF → DOCX conversion
from pdf2docx import Converter

import pandas as pd
import pdfplumber
import camelot
from pydantic import BaseModel, Field
from pypdf import PdfReader
from pdf2image import convert_from_path
from openpyxl import Workbook
from docx import Document as DocxWriter

# =========================================================================
# 1. ENVIRONMENT CONFIGURATION
# =========================================================================

load_dotenv()

gemini_key = os.environ.get("GEMINI_API_KEY")
groq_key = os.environ.get("GROQ_API_KEY")

# Use Gemma correctly (with the special header set in the client)
PRIMARY_TEXT = os.environ.get("PRIMARY_TEXT_MODEL", "gemma-4-26b-a4b-it")
BACKUP_TEXT = os.environ.get("BACKUP_TEXT_MODEL", "llama-3.3-70b-versatile")

PRIMARY_VISION = os.environ.get("PRIMARY_VISION_MODEL", "gemini-1.5-pro")
BACKUP_VISION = os.environ.get("BACKUP_VISION_MODEL", "gemini-1.5-flash")

if not gemini_key:
    print("❌ ERROR: Missing GEMINI_API_KEY in your .env file.")
    sys.exit(1)

if not groq_key:
    print("❌ ERROR: Missing GROQ_API_KEY in your .env file for backup redundancy.")
    sys.exit(1)

# [CHANGE] Enable Gemma models by adding the required HTTP header
google_client = genai.Client(
    api_key=gemini_key,
    http_options=types.HttpOptions(
        headers={"X-Goog-Enable-Gemma": "true"}
    )
)
groq_client = Groq(api_key=groq_key)

print(f"🔌 Core Engines Booted. Primary Text: {PRIMARY_TEXT} | Backup Text: Groq Ecosystem")
print("📄 Local PDF→DOCX Converter: pdf2docx (ready)")

# =========================================================================
# 2. DATA SCHEMAS (with description column)
# =========================================================================

class SpatialDimensionRow(BaseModel):
    feature_name: str = Field(description="Name of entry")
    length: str = Field(description="Length measurement with units")
    breadth: str = Field(description="Breadth measurement with units")
    height: str = Field(description="Height measurement with units")
    calculated_area: str = Field(description="Calculated area metric layout footprint")
    # [CHANGE] New field for image description
    description: str = Field(description="Brief 3-5 word description of the image or object", default="N/A")

class PureDimensionalResult(BaseModel):
    items: list[SpatialDimensionRow]

# =========================================================================
# 3. TEXT MODEL ROUTER (for summarisation)
# =========================================================================

def execute_text_with_fallback(prompt: str) -> str:
    """Sends a prompt to Gemini (with Gemma support), fails over to Groq if needed."""
    try:
        print(f"🧠 Sending summarisation request to {PRIMARY_TEXT}...")
        response = google_client.models.generate_content(
            model=PRIMARY_TEXT,
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.3,
                system_instruction="You are a professional document analyst. Provide concise, insightful summaries."
            )
        )
        return response.text.strip()
    except Exception as e:
        print(f"⚠️ Gemini failed: {e}")
        print(f"🔄 Falling back to Groq ({BACKUP_TEXT})...")
        try:
            chat_completion = groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": "You are a professional document analyst."},
                    {"role": "user", "content": prompt}
                ],
                model=BACKUP_TEXT,
                temperature=0.3,
            )
            return chat_completion.choices[0].message.content.strip()
        except Exception as groq_err:
            print(f"❌ Both text models failed: {groq_err}")
            return None

def generate_summary(text: str) -> str:
    """Generate a summary of the given text using the AI text model with strict constraints."""
    if not text or not text.strip():
        return "⚠️ No text content available for summarisation."

    prompt = (
        "You are given the full text of a document. Please write a comprehensive yet concise summary "
        "covering the main topics, key points, and any important conclusions. Use clear sections if needed.\n\n"
        "⚠️ STRICT CRITICAL FORMATTING RULES FOR WORD PROCESSING SUITABILITY:\n"
        "- DO NOT use LaTeX math notation, expressions, equations, or dollar signs ($).\n"
        "- Spell out all technical dimensions, Greek letters, variables, and units clearly in plain text English.\n"
        "- DO NOT embed any unparsed inline characters like '<3$%' or raw coding operators inside text paragraphs.\n"
        "- DO NOT use asterisks (**) inside sentences to bold isolated words.\n\n"
        f"DOCUMENT TEXT:\n{text}\n\nSUMMARY:"
    )
    return execute_text_with_fallback(prompt)

# =========================================================================
# 4. VISION CORE ROUTER (with different fallback model)
# =========================================================================

def execute_vision_with_fallback(model_name: str, contents_payload, config_payload):
    try:
        response = google_client.models.generate_content(
            model=model_name,
            contents=contents_payload,
            config=config_payload
        )
        return response
    except (google.genai.errors.APIError, Exception) as e:
        error_str = str(e)
        if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str or "quota" in error_str.lower():
            if model_name == PRIMARY_VISION:
                print(f"\n⚠️ [VISION RATE LIMIT] {PRIMARY_VISION} saturated. Dropping to fallback tier...")
                print(f"🔄 Routing to backup Vision model: {BACKUP_VISION}...")
                if BACKUP_VISION != PRIMARY_VISION:
                    return execute_vision_with_fallback(BACKUP_VISION, contents_payload, config_payload)
                else:
                    print("❌ Backup model is identical to primary; cannot bypass quota.")
                    raise e
        print(f"❌ Vision API Error: {e}")
        raise e

# =========================================================================
# 5. LOCAL PDF → DOCX EXTRACTION (pdf2docx) + AI SUMMARISATION
# =========================================================================

def extract_text_to_docx(pdf_path: str, docx_output_path: str):
    """
    Converts a digital PDF to a DOCX using pdf2docx.
    Returns True if successful.
    """
    try:
        print(f"⚙️  Converting PDF to DOCX with pdf2docx...")
        cv = Converter(pdf_path)
        cv.convert(docx_output_path)
        cv.close()
        print(f"✅ PDF converted successfully: {docx_output_path}")
        return True
    except Exception as e:
        print(f"❌ pdf2docx conversion failed: {e}")
        return False

def add_ai_summary_to_docx(docx_path: str, pdf_path: str = None):
    """
    Extracts text using a highly robust spacing-aware layout engine from the source PDF
    to fix merged spacing issues, requests an AI summary, cleans structural symbols,
    and appends it safely onto the document.
    """
    document_text = ""
    if pdf_path and os.path.exists(pdf_path):
        print("📖 Extracting spacing-aware layout text from PDF for AI processing matrix...")
        full_text_layers = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                p_text = page.extract_text(layout=True, x_tolerance=4, y_tolerance=3)
                if p_text:
                    full_text_layers.append(p_text)
        document_text = "\n".join(full_text_layers)
    
    if not document_text.strip():
        doc = DocxWriter(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        document_text = "\n".join(full_text)

    if not document_text.strip():
        print("⚠️ No text could be extracted – skipping AI summary.")
        return

    print("🧠 Generating AI summary of extracted text...")
    summary = generate_summary(document_text)
    if not summary:
        print("❌ Summary generation failed. DOCX remains unchanged.")
        return

    doc = DocxWriter(docx_path)
    doc.add_page_break()
    doc.add_heading("AI-Generated Document Summary", level=1)
    
    for section in summary.split("\n\n"):
        clean_section = section.strip()
        if not clean_section:
            continue
            
        if clean_section.startswith("###"):
            heading_text = clean_section.replace("###", "").replace("**", "").strip()
            doc.add_heading(heading_text, level=3)
        elif clean_section.startswith("##"):
            heading_text = clean_section.replace("##", "").replace("**", "").strip()
            doc.add_heading(heading_text, level=2)
        elif clean_section.startswith("#"):
            heading_text = clean_section.replace("#", "").replace("**", "").strip()
            doc.add_heading(heading_text, level=1)
        else:
            clean_section = re.sub(r'\*\*|__', '', clean_section)
            clean_section = clean_section.replace('$', '').replace('\\text', '').replace('\\mu', 'micro')
            doc.add_paragraph(clean_section)
            
    doc.save(docx_path)
    print(f"✨ AI summary appended to: {docx_path}")

# =========================================================================
# 6. IMAGE DETECTION AND EXTRACTION (multi‑page aware)
# =========================================================================

def check_pdf_has_images(pdf_path: str):
    """
    Returns a list of page numbers (1‑indexed) that contain at least one image.
    If no images, returns empty list. If detection fails, returns None to signal "convert all pages".
    """
    pages_with_images = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                if page.images:
                    print(f"✅ Found {len(page.images)} image(s) on page {page_num}")
                    pages_with_images.append(page_num)
        if pages_with_images:
            return pages_with_images
        else:
            print("ℹ️ No images detected in any page via pdfplumber")
            if _check_pdf_has_images_pypdf2(pdf_path):
                print("ℹ️ PyPDF2 indicates images exist but cannot determine pages – will convert all pages.")
                return None
            return []
    except Exception as e:
        print(f"⚠️ pdfplumber image check failed: {e}")
        if _check_pdf_has_images_pypdf2(pdf_path):
            print("ℹ️ PyPDF2 indicates images exist – will convert all pages.")
            return None
        return []

def _check_pdf_has_images_pypdf2(pdf_path: str) -> bool:
    try:
        reader = PdfReader(pdf_path)
        for page_num, page in enumerate(reader.pages):
            if "/Resources" in page and "/XObject" in page["/Resources"]:
                xobjects = page["/Resources"]["/XObject"].get_object()
                for obj in xobjects:
                    if xobjects[obj].get_object()["/Subtype"] == "/Image":
                        print(f"✅ PyPDF2 fallback: Found image on page {page_num+1}")
                        return True
        return False
    except Exception:
        return False

def extract_images_from_pdf(pdf_path: str, page_numbers=None, output_dir: str = None):
    """
    Extracts images from specified pages (1‑indexed).
    If page_numbers is None, converts all pages.
    Returns list of (image_bytes, width, height, page_num).
    """
    if output_dir is None:
        output_dir = os.path.dirname(pdf_path)

    poppler_path = os.environ.get("POPPLER_PATH", None)
    if not poppler_path or not os.path.exists(poppler_path):
        possible_paths = [
            r"D:\Poppler\Release-26.02.0-0\poppler-26.02.0\Library\bin",
            r"C:\Program Files\poppler\bin",
            r"C:\poppler\bin",
            "/usr/bin",
            "/usr/local/bin"
        ]
        for path in possible_paths:
            if os.path.exists(path):
                poppler_path = path
                break

    try:
        if page_numbers is None:
            first_page = 1
            last_page = None
        else:
            first_page = min(page_numbers)
            last_page = max(page_numbers)

        if poppler_path:
            print(f"📍 Using Poppler at: {poppler_path}")
            images = convert_from_path(pdf_path, first_page=first_page, last_page=last_page, poppler_path=poppler_path)
        else:
            print("ℹ️ Poppler not found, trying without explicit path...")
            images = convert_from_path(pdf_path, first_page=first_page, last_page=last_page)

        if not images:
            print("❌ No pages could be converted to images")
            return []

        image_data = []
        current_page = first_page
        for img in images:
            img_w, img_h = img.size
            print(f"📐 Page {current_page} dimensions: {img_w}x{img_h}px")

            working_img = img.copy()
            if img_w > 2048 or img_h > 2048:
                working_img.thumbnail((2048, 2048))
                print(f"   Resized to: {working_img.size}")

            img_byte_arr = io.BytesIO()
            working_img.save(img_byte_arr, format='JPEG', quality=90)
            image_data.append((img_byte_arr.getvalue(), img_w, img_h, current_page))
            current_page += 1

        return image_data

    except Exception as e:
        print(f"❌ Error extracting images from PDF: {e}")
        return []

# =========================================================================
# 6b. VISION ANALYSIS WITH DESCRIPTION (updated)
# =========================================================================

def analyze_canvas_dimensions_native(img_bytes, img_w, img_h):
    # [CHANGE] Added instruction to provide a short description
    prompt_base = (
        f"The layout image boundaries are exactly {img_w} wide by {img_h} high in pixels.\n"
        "Extract numerical dimensions from the attached page layout view and provide a short description of the image.\n\n"
        "STRICT DATA REGISTRATION RULES:\n"
        "1. FIRST ENTRY: Always register the total overall image canvas layout footprint under feature_name='Overall Image Canvas'.\n"
        "2. INTERNAL SYMMETRIC OBJECTS: If there are distinct symmetrical geometric shapes located inside the image boundary, map their specific structural boundaries as additional rows.\n"
        "3. DESCRIPTION: For the 'Overall Image Canvas' and each major object, provide a brief 3-5 word description (e.g., 'Site plan', 'Mechanical part', 'Circuit diagram').\n"
        "4. POPULATE: Length, Breadth, Height, Area (with units), and Description for each tracked item."
    )

    try:
        print("🛰️ Streaming layout matrix to Vision Engine [Mode A: Schema Extraction]...")
        response = execute_vision_with_fallback(
            model_name=PRIMARY_VISION,
            contents_payload=[
                types.Part.from_bytes(data=img_bytes, mime_type='image/jpeg'),
                prompt_base + "\nStrictly adhere to the requested structural JSON response schema."
            ],
            config_payload=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=PureDimensionalResult,
                temperature=0.0
            )
        )
        if response.parsed and hasattr(response.parsed, 'items'):
            items = response.parsed.items
            if items and len(items) > 0:
                print(f"✅ Mode A successful: Found {len(items)} spatial items with descriptions")
                return items
    except Exception as e:
        print(f"⚠️ Mode A hit an operational fault: {e}. Moving to Mode B fallback...")

    # -----------------------------------------------------------------
    # MODE B: Text Fallback (now with 6 columns including description)
    # -----------------------------------------------------------------
    print("🛰️ Streaming layout matrix to Vision Engine [Mode B: Text Parsing Fallback]...")
    fallback_prompt = (
        prompt_base + "\n"
        "Output your findings as a clean 6-column table, separated with a vertical bar '|'.\n"
        "Columns: Feature Name | Length | Breadth | Height | Area | Description\n"
        "Example format:\n"
        "Overall Image Canvas | 1200px | 800px | 0px | 960000 sq px | Site Layout\n"
        "Internal Box | 400px | 400px | 0px | 160000 sq px | Inner component"
    )

    try:
        response = execute_vision_with_fallback(
            model_name=PRIMARY_VISION,
            contents_payload=[
                types.Part.from_bytes(data=img_bytes, mime_type='image/jpeg'),
                fallback_prompt
            ],
            config_payload=types.GenerateContentConfig(temperature=0.0)
        )

        text_lines = response.text.strip().split('\n')
        parsed_items = []
        for line in text_lines:
            if "|" in line and "Feature Name" not in line and "---" not in line:
                parts = [p.strip() for p in line.split("|")]
                if len(parts) >= 6:
                    # We have 6 columns
                    f, l, b, h, area, desc = parts[0], parts[1], parts[2], parts[3], parts[4], parts[5]
                elif len(parts) >= 5:
                    # Fallback if description missing – set as "N/A"
                    f, l, b, h, area = parts[0], parts[1], parts[2], parts[3], parts[4]
                    desc = "N/A"
                else:
                    continue

                class RowObject:
                    def __init__(self, f, l, b, h, a, d):
                        self.feature_name = f
                        self.length = l
                        self.breadth = b
                        self.height = h
                        self.calculated_area = a
                        self.description = d
                parsed_items.append(RowObject(f, l, b, h, area, desc))

        if parsed_items:
            print(f"✅ Mode B successful: Found {len(parsed_items)} spatial items with descriptions")
            return parsed_items
    except Exception as e:
        print(f"❌ Mode B also failed: {e}")

    return None

# =========================================================================
# 7. MASTER PIPELINE ENGINE ORCHESTRATOR
# =========================================================================

def run_comprehensive_individual_pipeline(pdf_path: str):
    local_script_folder = os.path.dirname(os.path.abspath(__file__)) if '__file__' in locals() else os.getcwd()
    file_base = os.path.splitext(os.path.basename(pdf_path))[0]

    text_output_file = os.path.join(local_script_folder, f"{file_base}_Extracted_Text.docx")
    tables_output_file = os.path.join(local_script_folder, f"{file_base}_Extracted_Tables.xlsx")
    vision_output_file = os.path.join(local_script_folder, f"{file_base}_Image_Dimensions.xlsx")

    # ---------------------------------------------------------------------
    # PHASE 1: PDF → DOCX + AI Summary
    # ---------------------------------------------------------------------
    if os.path.exists(text_output_file):
        print(f"ℹ️ Skipping PHASE 1: '{os.path.basename(text_output_file)}' already exists.")
    else:
        print("⚡ [PHASE 1] Converting PDF to DOCX locally with pdf2docx...")
        success = extract_text_to_docx(pdf_path, text_output_file)
        if success:
            print("📄 Adding AI-generated summary to the DOCX...")
            add_ai_summary_to_docx(text_output_file, pdf_path)
        else:
            print("❌ PHASE 1 failed. No DOCX created.")

    # ---------------------------------------------------------------------
    # PHASE 2: Table Isolation Engine (unchanged)
    # ---------------------------------------------------------------------
    if os.path.exists(tables_output_file):
        print(f"ℹ️ Skipping PHASE 2: '{os.path.basename(tables_output_file)}' already exists.")
    else:
        print("\n⚡ [PHASE 2] Initializing Table Isolation Processing...")
        valid_tables_found = {}
        table_counter = 0

        try:
            print("📊 Checking layouts via Camelot Engine...")
            for flavor in ["stream", "lattice"]:
                try:
                    camelot_tables = camelot.read_pdf(pdf_path, pages="all", flavor=flavor)
                    for tbl in camelot_tables:
                        df = tbl.df
                        if df.empty or tbl.parsing_report.get("accuracy", 0) < 80 or len(df) < 2 or len(df.columns) < 3:
                            continue

                        non_empty_cols = sum(1 for col in df.columns if (df[col].astype(str).str.strip() != "").sum() > 0)
                        if non_empty_cols < 3:
                            continue

                        total_cells = df.size
                        text_heavy_cells = sum(df[col].astype(str).str.count(r'\b\w+\b').gt(4).sum() for col in df.columns)

                        if total_cells > 0 and (text_heavy_cells / total_cells) > 0.15:
                            print(f"⏩ Dropped document-layout prose leak (Density: {text_heavy_cells/total_cells:.1%})")
                            continue

                        df = df.replace(r'\n+', ' ', regex=True).replace(r'\s+', ' ', regex=True)
                        table_counter += 1
                        valid_tables_found[f"Camelot_{flavor}_Table_{table_counter}"] = df.values.tolist()
                except Exception as e:
                    print(f"⚠️ Camelot {flavor} parser bypass triggered: {e}")
        except Exception as e:
            print(f"⚠️ Camelot overall error: {e}")

        if len(valid_tables_found) == 0:
            print("🔄 Activating PDFPlumber extraction fallback...")
            table_settings = {"vertical_strategy": "text", "horizontal_strategy": "text", "snap_tolerance": 6, "join_tolerance": 6, "intersection_tolerance": 8}
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        for tbl in page.find_tables(table_settings=table_settings):
                            extracted = tbl.extract()
                            if not extracted or len(extracted) < 2:
                                continue
                            table_counter += 1
                            cleaned_rows = [[re.sub(r"\s+", " ", str(cell or "")).strip() for cell in row] for row in extracted]
                            valid_tables_found[f"Page{page_num+1}_Table{table_counter}"] = cleaned_rows
                    except Exception as page_error:
                        print(f"⚠️ Plumber page step fault: {page_error}")

        if len(valid_tables_found) > 0:
            from openpyxl.utils import get_column_letter
            from openpyxl.styles import Alignment

            wb_tables = Workbook()
            default_sheet = wb_tables.active
            wb_tables.remove(default_sheet)

            for sheet_name, table_rows in valid_tables_found.items():
                ws = wb_tables.create_sheet(title=sheet_name[:31])
                ws.views.sheetView[0].showGridLines = True
                for row in table_rows:
                    ws.append(row)

                for col in ws.columns:
                    max_len = 0
                    col_letter = get_column_letter(col[0].column)
                    for cell in col:
                        if cell.value:
                            val_str = str(cell.value)
                            max_len = max(max_len, len(val_str))
                            if len(val_str) > 20:
                                cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
                            else:
                                cell.alignment = Alignment(vertical="center")
                    ws.column_dimensions[col_letter].width = min(max(max_len + 3, 11), 45)

            wb_tables.save(tables_output_file)
            print(f"✨ SUCCESS! Compiled {len(valid_tables_found)} pristine structured sheet matrices.")
            print(f"📁 Target Destination: {tables_output_file}")
        else:
            print("ℹ️ No operational tables recovered from document.")

    # ---------------------------------------------------------------------
    # PHASE 3: Canvas Spatial Vision Mapping Engine (with description column)
    # ---------------------------------------------------------------------
    if os.path.exists(vision_output_file):
        print(f"ℹ️ Skipping PHASE 3: '{os.path.basename(vision_output_file)}' already exists.")
    else:
        print("\n⚡ [PHASE 3] Scanning for images and extracting dimensions + descriptions...")
        pages_with_images = check_pdf_has_images(pdf_path)
        
        if pages_with_images is None:
            print("ℹ️ Converting all pages to images (fallback).")
            image_data_list = extract_images_from_pdf(pdf_path, page_numbers=None)
        elif len(pages_with_images) == 0:
            print("ℹ️ No images detected in PDF. Creating empty dimension file...")
            wb_vision = Workbook()
            ws_vision = wb_vision.active
            ws_vision.title = "No Images Found"
            ws_vision.append(["Note: No images were detected in this PDF document"])
            wb_vision.save(vision_output_file)
            image_data_list = []
        else:
            print(f"📸 Images detected on pages: {pages_with_images}")
            image_data_list = extract_images_from_pdf(pdf_path, page_numbers=pages_with_images)

        if not image_data_list:
            if not os.path.exists(vision_output_file):
                wb_vision = Workbook()
                ws_vision = wb_vision.active
                ws_vision.title = "No Images Found"
                ws_vision.append(["Note: No images could be extracted or no images found."])
                wb_vision.save(vision_output_file)
            print("ℹ️ No image data to process.")
        else:
            all_vision_rows = []
            for img_bytes, img_w, img_h, page_num in image_data_list:
                print(f"\n🔍 Analyzing page {page_num} image ({img_w}x{img_h}px)...")
                vision_rows = analyze_canvas_dimensions_native(img_bytes, img_w, img_h)

                if vision_rows:
                    all_vision_rows.extend(vision_rows)
                    print(f"   ✅ Found {len(vision_rows)} spatial items with descriptions")
                else:
                    print(f"   ⚠️ No dimensions extracted from page {page_num}")

            if all_vision_rows:
                print(f"\n📊 Creating dimension workbook with {len(all_vision_rows)} total entries...")
                wb_vision = Workbook()
                ws_vision = wb_vision.active
                ws_vision.title = "Dimensions Ledger"
                ws_vision.views.sheetView[0].showGridLines = True
                # [CHANGE] Updated header to include Description column
                ws_vision.append(["Feature Name", "Length", "Breadth", "Height", "Estimated Area", "Image/Object Description"])

                for item in all_vision_rows:
                    # [CHANGE] Append the description as the 6th column
                    ws_vision.append([
                        item.feature_name,
                        item.length,
                        item.breadth,
                        item.height,
                        item.calculated_area,
                        getattr(item, 'description', 'N/A')  # fallback if missing
                    ])

                for col in ws_vision.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws_vision.column_dimensions[column].width = adjusted_width

                wb_vision.save(vision_output_file)
                print(f"✨ SUCCESS! Graphic metric sheet compiled: {vision_output_file}")
            else:
                print("⚠️ No dimensions could be extracted from any images")
                if not os.path.exists(vision_output_file):
                    wb_vision = Workbook()
                    ws_vision = wb_vision.active
                    ws_vision.title = "No Dimensions"
                    ws_vision.append(["Note: Images were found but no dimensions could be extracted"])
                    wb_vision.save(vision_output_file)

    print("\n🏁 [COMPLETE] Core processing worker finished run cycle.")
    print(f"📁 Output files:")
    print(f"   - Text + Summary (DOCX): {os.path.basename(text_output_file)}")
    print(f"   - Tables (XLSX): {os.path.basename(tables_output_file)}")
    print(f"   - Dimensions (XLSX): {os.path.basename(vision_output_file)}")


if __name__ == "__main__":
    print("====================================================")
    print("      CAiVision Advanced Hybrid Document Pipeline  ")
    print("====================================================\n")

    while True:
        user_input = input("📥 Enter the name of your PDF file (e.g., sample_data.pdf): ").strip()
        if user_input and not user_input.lower().endswith('.pdf'):
            user_input += '.pdf'

        if os.path.exists(user_input):
            target_pdf = user_input
            print(f"✅ Found target file: '{target_pdf}'")
            break
        else:
            print(f"❌ Error: The file '{user_input}' could not be found in this directory.\n")

    run_comprehensive_individual_pipeline(target_pdf)