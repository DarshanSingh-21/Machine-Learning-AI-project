import os
import io
import sys
import time
import json
import re
from dotenv import load_dotenv
from google import genai
from google.genai import types
import google.genai.errors
import pandas as pd
import pdfplumber
from pydantic import BaseModel, Field
from pdf2image import convert_from_path
from PIL import Image
from openpyxl import Workbook
from docx import Document as DocxWriter

# =========================================================================
# 1. ENVIRONMENT CONFIGURATION & GOOGLE GENAI SDK BOOT
# =========================================================================

load_dotenv()

active_api_key = os.environ.get("API_KEY") or os.environ.get("GEMINI_API_KEY")

if not active_api_key:
    print("❌ ERROR: No valid API key discovered in your environment or .env file.")
    sys.exit(1)

client = genai.Client(api_key=active_api_key)
print("🔌 System Engine Booted successfully using Native Google GenAI SDK Client.")

# =========================================================================
# 2. PURE METRICS-ONLY DATA SCHEMA FOR IMAGE ELEMENT
# =========================================================================

class SpatialDimensionRow(BaseModel):
    feature_name: str = Field(description="Name of entry")
    length: str = Field(description="Length measurement with units")
    breadth: str = Field(description="Breadth measurement with units")
    height: str = Field(description="Height measurement with units")
    calculated_area: str = Field(description="Calculated area metric layout footprint")

class PureDimensionalResult(BaseModel):
    items: list[SpatialDimensionRow]

# =========================================================================
# 3. SMART RATE-LIMIT HANDLING ENGINE (BACKOFF ALGORITHM)
# =========================================================================

def execute_with_quota_retry(api_call_func, *args, **kwargs):
    max_attempts = 4
    base_delay = 15
    
    for attempt in range(max_attempts):
        try:
            return api_call_func(*args, **kwargs)
        except google.genai.errors.APIError as e:
            if e.code == 429 or "RESOURCE_EXHAUSTED" in str(e):
                print(f"\n⚠️ [QUOTA EXHAUSTED] Hit Gemini limit ceiling.")
                err_msg = str(e)
                wait_match = re.search(r"retry in ([\d\.]+)\s*s", err_msg, re.IGNORECASE) or \
                             re.search(r"retry in ([\d\.]+)\s*ms", err_msg, re.IGNORECASE)
                
                sleep_duration = base_delay
                if wait_match:
                    parsed_time = float(wait_match.group(1))
                    if "ms" in err_msg.lower() and "retry in 0s" not in err_msg.lower():
                        sleep_duration = parsed_time / 1000.0
                    else:
                        sleep_duration = parsed_time
                
                sleep_duration = max(int(sleep_duration) + 3, 5)
                print(f"⏳ Sleeping for {sleep_duration} seconds before retry ({attempt + 1}/{max_attempts})...")
                time.sleep(sleep_duration)
            else:
                raise e
        except Exception as e:
            if "429" in str(e) or "quota" in str(e).lower():
                print(f"\n⚠️ [QUOTA EXHAUSTED] Sleeping 65 seconds...")
                time.sleep(65)
            else:
                raise e
    print("❌ Error: Maximum structural backoff retry attempts exhausted.")
    return None

# =========================================================================
# 4. EXTRACTION WORKER SUB-FUNCTIONS
# =========================================================================

def extract_raw_pdf_text(pdf_path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(pdf_path)
    extracted_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            extracted_text.append(text)
    return "\n".join(extracted_text).strip()

def save_analysis_to_docx(analysis_content: str, docx_output_path: str, report_title: str):
    doc = DocxWriter()
    doc.add_heading(report_title, level=0)
    paragraphs = analysis_content.split('\n')
    for p in paragraphs:
        clean_p = p.strip()
        if clean_p.startswith("###"):
            doc.add_heading(clean_p.replace("###", "").strip(), level=3)
        elif clean_p.startswith("##"):
            doc.add_heading(clean_p.replace("##", "").strip(), level=2)
        elif clean_p.startswith("#"):
            doc.add_heading(clean_p.replace("#", "").strip(), level=1)
        elif clean_p:
            doc.add_paragraph(clean_p)
    doc.save(docx_output_path)
    print(f"✨ Text report written to Word document: {docx_output_path}")

def run_text_analysis_layer(raw_text: str) -> str:
    print("🧠 Processing Text Intelligence Layer...")
    prompt = (
        "Perform a rigorous thematic analysis of the provided text. Do not summarize briefly.\n"
        "1. Generate a comprehensive executive summary outlining overarching objectives.\n"
        "2. Isolate the top 5 most critical technical trends, system metrics, or structural data insights.\n"
        "3. Elaborate deeply on each insight—provide domain context and actionable next steps.\n\n"
        f"Raw Extracted Text Context:\n{raw_text}"
    )
    
    def api_wrapper():
        response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        return response.text.strip()
        
    return execute_with_quota_retry(api_wrapper)

def analyze_canvas_dimensions_native(img_bytes, img_w, img_h):
    prompt_base = (
        f"The layout image boundaries are exactly {img_w} wide by {img_h} high in pixels.\n"
        "Extract ONLY numerical dimensions from the attached page layout view. DO NOT provide any text descriptions, notes, or explanations.\n\n"
        "STRICT DATA REGISTRATION RULES:\n"
        "1. FIRST ENTRY: Always register the total overall image canvas layout footprint under feature_name='Overall Image Canvas'.\n"
        "2. INTERNAL SYMMETRIC OBJECTS: If there are distinct symmetrical geometric shapes located inside the image boundary, map their specific structural boundaries as additional rows.\n"
        "3. NO TEXT/DESCRIPTIONS: Populate ONLY raw measurements (Length, Breadth, Height, Area with units) for each tracked item.\n"
        "If you cannot determine physical units (like cm or mm), use pixel counts (px) scaled against the canvas boundaries."
    )

    # --- ACTION 1: STRUCTURED SCHEMA ATTEMPT ---
    def schema_call():
        print(f"🛰️ Streaming layout matrix to Vision Engine [Mode A: Schema Extraction]...")
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                types.Part.from_bytes(data=img_bytes, mime_type='image/jpeg'),
                prompt_base + "\nStrictly adhere to the requested structural JSON response schema."
            ],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=PureDimensionalResult,
                temperature=0.0
            ),
        )
        return response.parsed.items if (response.parsed and hasattr(response.parsed, 'items')) else None

    try:
        result = execute_with_quota_retry(schema_call)
        if result:
            return result
    except Exception as e:
        print(f"⚠️ Mode A hit an operational fault: {e}. Moving to Mode B fallback layout directly...")

    # --- ACTION 2: MARKDOWN TABLE PARSING FALLBACK ATTEMPT ---
    def fallback_call():
        print(f"🛰️ Streaming layout matrix to Vision Engine [Mode B: Text Parsing Fallback]...")
        fallback_prompt = (
            prompt_base + "\n"
            "Output your findings as a clean table matching this style, separation with a vertical bar '|'.\n"
            "Example format:\n"
            "Overall Image Canvas | 1200px | 800px | 0px | 960000 sq px\n"
            "Internal Box | 400px | 400px | 0px | 160000 sq px"
        )
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                types.Part.from_bytes(data=img_bytes, mime_type='image/jpeg'),
                fallback_prompt
            ],
            config=types.GenerateContentConfig(temperature=0.0),
        )
        
        text_lines = response.text.strip().split('\n')
        parsed_items = []
        for line in text_lines:
            if "|" in line and "Feature Name" not in line and "---" not in line:
                parts = [p.strip() for p in line.split("|")]
                if len(parts) >= 4:
                    while len(parts) < 5:
                        parts.append("0")
                    # Fixed Class properties configuration initialization mapping layout errors
                    class RowObject:
                        def __init__(self, f, l, b, h, a):
                            self.feature_name = f
                            self.length = l
                            self.breadth = b
                            self.height = h
                            self.calculated_area = a
                    parsed_items.append(RowObject(parts[0], parts[1], parts[2], parts[3], parts[4]))
        return parsed_items if len(parsed_items) > 0 else None

    return execute_with_quota_retry(fallback_call)

# =========================================================================
# 5. MASTER COMPREHENSIVE PIPELINE ORCHESTRATOR
# =========================================================================

def run_comprehensive_individual_pipeline(pdf_path: str):
    if not os.path.exists(pdf_path):
        print(f"❌ Error: Target file path '{pdf_path}' does not exist.")
        return

    local_script_folder = os.path.dirname(os.path.abspath(__file__)) if '__file__' in locals() else os.getcwd()
    file_base = os.path.splitext(os.path.basename(pdf_path))[0]
    
    print(f"📂 Output target directory: {local_script_folder}\n")

    text_output_file = os.path.join(local_script_folder, f"{file_base}_Text_Deep_Analysis.docx")
    tables_output_file = os.path.join(local_script_folder, f"{file_base}_Extracted_Tables.xlsx")
    vision_output_file = os.path.join(local_script_folder, f"{file_base}_Image_Dimensions.xlsx")

    # ---------------------------------------------------------------------
    # FILE 1: Text Intelligence Generation (.docx)
    # ---------------------------------------------------------------------
    if os.path.exists(text_output_file):
        print(f"ℹ️ Skipping PHASE 1: '{os.path.basename(text_output_file)}' already exists.")
    else:
        print("⚡ [PHASE 1] Scanning structural text content layers...")
        raw_document_text = extract_raw_pdf_text(pdf_path)
        if raw_document_text:
            text_analysis_content = run_text_analysis_layer(raw_document_text)
            if text_analysis_content:
                save_analysis_to_docx(text_analysis_content, text_output_file, "CAiVision Advanced Document Text Analysis")
        else:
            print("ℹ️ No textual layer resolved; skipping text document creation.")

    # ---------------------------------------------------------------------
    # FILE 2: ADVANCED DATA GRID EXTRACTION CRITERIA (.xlsx)
    # ---------------------------------------------------------------------
    if os.path.exists(tables_output_file):
        print(f"ℹ️ Skipping PHASE 2: '{os.path.basename(tables_output_file)}' already exists.")
    else:
        print("\n⚡ [PHASE 2] Auditing document coordinate layout with explicit table algorithms...")
        wb_tables = Workbook()
        default_sheet = wb_tables.active
        wb_tables.remove(default_sheet)
        
        table_extraction_settings = {
            "vertical_strategy": "text", "horizontal_strategy": "text",     
            "snap_tolerance": 3, "join_tolerance": 3, "edge_min_length": 15              
        }
        
        table_counter = 0
        with pdfplumber.open(pdf_path) as pdf:
            for idx, page in enumerate(pdf.pages):
                extracted_tables = page.extract_tables(table_settings=table_extraction_settings)
                for t_idx, current_table in enumerate(extracted_tables):
                    table_counter += 1
                    ws_table = wb_tables.create_sheet(title=f"Page{idx+1}_Table{t_idx+1}")
                    ws_table.views.sheetView[0].showGridLines = True
                    for data_row in current_table:
                        clean_row = ["" if val is None else str(val).strip() for val in data_row]
                        ws_table.append(clean_row)
                        
        if table_counter > 0:
            wb_tables.save(tables_output_file)
            print(f"✨ SUCCESS! Extracted {table_counter} distinct tables: {tables_output_file}")
        else:
            print("ℹ️ No multi-column table matrix identified using fallback settings.")

    # ---------------------------------------------------------------------
    # FILE 3: Canvas Vision Engine Layer Image Specs Workbook (.xlsx)
    # ---------------------------------------------------------------------
    if os.path.exists(vision_output_file):
        print(f"ℹ️ Skipping PHASE 3: '{os.path.basename(vision_output_file)}' already exists.")
    else:
        print("\n⚡ [PHASE 3] Rendering canvas page matrices into raster assets...")
        try:
            poppler_bin_path = r"D:\Poppler\Release-26.02.0-0\poppler-26.02.0\Library\bin"

            if os.path.exists(poppler_bin_path):
                pages = convert_from_path(pdf_path, first_page=1, last_page=1, poppler_path=poppler_bin_path)
            else:
                print(f"❌ Critical Error: Poppler path invalid at: '{poppler_bin_path}'")
                return

            if pages:
                page_image = pages[0].convert("RGB")
                img_w, img_h = page_image.size
                print(f"📐 Canvas Boundary Bounds Discovered: {img_w}x{img_h}px")
                
                # Resizing copy locally to prevent modification anomalies inside memory
                working_img = page_image.copy()
                working_img.thumbnail((1024, 1024))
                
                img_byte_arr = io.BytesIO()
                working_img.save(img_byte_arr, format='JPEG', quality=85)
                compressed_bytes = img_byte_arr.getvalue()
                
                vision_rows = analyze_canvas_dimensions_native(compressed_bytes, img_w, img_h)
                
                if vision_rows:
                    wb_vision = Workbook()
                    ws_vision = wb_vision.active
                    ws_vision.title = "Dimensions Ledger"
                    ws_vision.views.sheetView[0].showGridLines = True
                    ws_vision.append(["Feature Name", "Length", "Breadth", "Height", "Estimated Area"])
                    
                    for item in vision_rows:
                        ws_vision.append([
                            item.feature_name, item.length, item.breadth, item.height, item.calculated_area
                        ])
                    
                    wb_vision.save(vision_output_file)
                    print(f"✨ SUCCESS! Standalone Image metrics sheet compiled: {vision_output_file}")
                else:
                    print("❌ Error: Vision Extraction completely exhausted all backoff options.")
        except Exception as e:
            print(f"⚠️ Vision synthesis worker encountered an obstacle: {e}")

    print("\n🏁 [SUCCESS] Pipeline execution evaluated completely.")

if __name__ == "__main__":
    target_pdf = "sample_data.pdf" 
    run_comprehensive_individual_pipeline(target_pdf)